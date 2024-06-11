import datetime
import io
import os
import pytz
import gspread
import numpy as np
import pandas as pd
import streamlit as st
from docx import Document
from numpy import datetime_as_string


utc_now = pytz.utc.localize(datetime.datetime.utcnow())
t_pytz = utc_now.astimezone(pytz.timezone("US/Central"))
# extract year, month and day from t_pytz and store under variable local_year, local_month and local_day
local_year = t_pytz.year
local_month = t_pytz.month
local_day = t_pytz.day

# extract hour and minute from t_pytz and store under variable local_hour and local_minute
local_hour = t_pytz.hour
local_minute = t_pytz.minute

st.session_state.radio_dict = {}
if st.session_state.radio_dict is None:
    st.session_state.radio_dict = {}

# Initialize or load existing data from CSV
try:
    df = pd.read_csv('rvu.csv')
except FileNotFoundError:
    df = pd.DataFrame(columns=['Date', 'Time', 'CPT', 'RVU'])

#google_sheet_url = st.secrets["private_gsheets_url"]

def convert_datatime_to_string(date, time):
    """
    Convert datetime to a date string and time string
    Then reurn the date string and time string
    """
    date_string = str(date)
    date_time_string = str(time)
    return date_string, date_time_string


def save_into_csv(date,time, cpt, wrvu):
    """
    Save the data into a  csv file
    """
    date, time_stamp = convert_datatime_to_string(date, time)
    data = {'Date': [date], 'Time': [time_stamp], 'CPT': [cpt], 'RVU': [wrvu]}
    df_new = pd.DataFrame(data)
    df = pd.concat([df, df_new], ignore_index=True)
    df.to_csv('rvu.csv', index=False)


# Function to get the value of the selected radio button and return the value
def get_value(key):
    value = st.session_state[key]
    return value


# function that returns value from the dictionary based on the key
def get_rvu_value(key):
    value = st.session_state.radio_dict[key]
    return value


def get_label():
    keys = [
        "opfollowup",
        "opnewpatient",
        "procedure",
        "epicconsult",
        "ipfollowup",
        "ipnewpatient",
    ]
    for key in keys:
        value = get_value(key)
        if value != "None":
            return value


# A function to read all the data from the google sheet and return it as a dataframe
def read_data_from_google_sheet():
    try:
        df = pd.read_csv('rvu.csv')
    except FileNotFoundError:
        df = pd.DataFrame(columns=['Date', 'Time', 'CPT', 'RVU'])

    return df


# A function caled report that will generate a report based on the data in the google sheet.
# The function will take a dataframe as an input and return total wrvu, it will also return a dataframe with the total wrvu for cpt codes
def report(df):
    df["wrvu"] = pd.to_numeric(df["wrvu"])
    total_wrvu = df["wrvu"].sum()
    df = df.groupby("cpt").sum("wrvu")
    return total_wrvu, df


# A function caled report that will generate a report based on the data in the google sheet.
# The function will take a dataframe as an input and two dates and return total wrvu, it will also return a dataframe with the total wrvu for cpt codes within the data range and count for each cpt code
def report_by_date(df, start_date, end_date):
    df["wrvu"] = pd.to_numeric(df["wrvu"])
    df["date"] = pd.to_datetime(df["date"])
    # convert date column from datetime64[ns] to datetime otherwise comparison will not work
    df["date"] = df["date"].dt.date
    df = df[(df["date"] >= start_date) & (df["date"] <= end_date)]
    total_wrvu = df["wrvu"].sum()
    df = df.groupby("cpt").agg({"Count": "count", "wrvu": "sum"})
    return total_wrvu, df


# A function to generate a report when the 'generate report' button is clicked
def generate_report():
    df = read_data_from_google_sheet()
    total_wrvu, wrvu_by_cpt = report_by_date(df, start_date, end_date)
    # format wrvu_by_cpt to 2 decimal places
    total_wrvu = total_wrvu.round(2)
    st.sidebar.header("Total wrvu: ")
    st.sidebar.header(total_wrvu)
    st.sidebar.header("wRVU by CPT")
    st.sidebar.dataframe(wrvu_by_cpt)


# A function to generate a word document with total_wrvu and wrvu_by_cpt when the 'Downlad report' button is clicked
def download_report():
    df = read_data_from_google_sheet()
    total_wrvu, wrvu_by_cpt = report_by_date(df, start_date, end_date)
    # format wrvu_by_cpt to 2 decimal places
    total_wrvu = total_wrvu.round(2)
    wrvu_by_cpt = wrvu_by_cpt.round(2)
    # create a word document with total_wrvu and wrvu_by_cpt
    document = Document()
    document.add_heading("wRVU Report", 0)
    # Add date range used to generate the report
    document.add_paragraph(
        "Date range for the report: from " + str(start_date) + "   to   " + str(end_date)
    )
    document.add_heading("Total wrvu: ", 0)
    document.add_paragraph(str(total_wrvu))
    document.add_heading("wRVU by CPT", 0)
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "CPT"
    hdr_cells[1].text = "Count"
    hdr_cells[2].text = "wRVU"
    for index, row in wrvu_by_cpt.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = index
        row_cells[1].text = str(row["Count"])
        row_cells[2].text = str(row["wrvu"])
    # document.save("report.docx") - this can be used to save the document to the local machine
    # add a footer to the document stating "Created by wRVU monitor app"
    section = document.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.text = "Created by wRVU monitor app"
    return document


# create python dictionary to store the values of the radio buttons


st.session_state.radio_dict = {
    "None": 0,
    "Op Fu Level 1 99211": 0.18,
    "Op Fu Level 2 99212": 0.7,
    "Op Fu Level 3 99213": 1.3,
    "Op Fu Level 4 99214": 1.92,
    "Op Fu Level 5 99215": 2.8,
    "Op new Level 1 99201": 0,
    "Op new Level 2 99202": 0.93,
    "Op new Level 3 99203": 1.6,
    "Op new Level 4 99204": 2.6,
    "Op new Level 5 99205": 3.5,
    "1_FNA": 1.46,
    "2_FNA": 2.46,
    "CGM_read": 0.7,
    "US 76536": 0.56,
    "5 min or more 99451": 0.7,
    "5-10 min 99446": 0.35,
    "11-20 min 99447": 0.7,
    "21-30 min 99448": 1.05,
    "31 min or more 99449": 1.4,
    "Ip fu Level 2 99232": 1.39,
    "Ip fu Level 3 99233": 2,
    "Ip new Level 2 99222": 2.61,
    "Ip new Level 3 99223": 3.86,
}

st.title("RVU App")

dateco, timecol = st.columns(2)
with dateco:
    date = st.date_input(
        "Date of the visit",
        datetime.date(local_year,local_month, local_day)  # type: ignore
    )

with timecol:
    time = st.time_input(
        "Time of the visit",
        datetime.time(local_hour, local_minute),  # type: ignore
    )

with st.form("my_form", clear_on_submit=True):
    st.write("Out patient visits")
    (
        ucol1,
        ucol2,
    ) = st.columns(2)

    with ucol1:
        st.radio(
            "Follow up visits",
            [
                "None",
                "Op Fu Level 1 99211",
                "Op Fu Level 2 99212",
                "Op Fu Level 3 99213",
                "Op Fu Level 4 99214",
                "Op Fu Level 5 99215",
            ],
            key="opfollowup",
        )

    with ucol2:
        st.radio(
            "New Patient visits",
            [
                "None",
                "Op new Level 1 99201",
                "Op new Level 2 99202",
                "Op new Level 3 99203",
                "Op new Level 4 99204",
                "Op new Level 5 99205",
            ],
            key="opnewpatient",
        )
    st.write("Procedures and misc")
    (
        pcol1,
        pcol2,
    ) = st.columns(2)
    with pcol1:
        st.radio(
            "FNA",
            ["None", "1_FNA", "2_FNA", "CGM_read", "US 76536"],
            key="procedure",
        )
    with pcol2:
        st.radio(
            "Interprofessional tele Consults",
            [
                "None",
                "5 min or more 99451",
                "5-10 min 99446",
                "11-20 min 99447",
                "21-30 min 99448",
                "31 min or more 99449",
            ],
            key="epicconsult",
        )
    st.write("Inpatient visits")

    (
        icol1,
        icol2,
    ) = st.columns(2)
    with icol1:
        st.radio(
            "Follow up visits",
            ["None", "Ip fu Level 2 99232", "Ip fu Level 3 99233"],
            key="ipfollowup",
        )
    with icol2:
        st.radio(
            "New Patient visits",
            ["None", "Ip new Level 2 99222", "Ip new Level 3 99223"],
            key="ipnewpatient",
        )

    submitted = st.form_submit_button("Submit")
    if submitted:
        st.write("You selected:", get_label())
        cpt_code = get_label()
        rvu_value = get_rvu_value(cpt_code)
        save_into_csv(date,time, cpt_code, rvu_value)
        st.write("Data saved into Google Sheet")

# Sidebar
# Create Streamlit sidebar with a button named Generate Report and download the report

st.sidebar.title("RVU Report")
st.sidebar.markdown("Click the button below to generate the report")
start_date = st.sidebar.date_input("Start Date",datetime.date(local_year,local_month, local_day), key="start_date")
end_date = st.sidebar.date_input("End Date", datetime.date(local_year,local_month, local_day),key="end_date")
gen_button = st.sidebar.button("Generate Report", key="generate_report")

if gen_button:
    generate_report()

with st.sidebar:
    doc_download = download_report()
    bio = io.BytesIO()
    doc_download.save(bio)
    if doc_download:
        st.download_button(
            label="Download Report",
            data=bio.getvalue(),
            file_name="wRVU_Report_"
            + str(np.datetime64(datetime.datetime.now()))
            + ".docx",
            mime="docx",
        )